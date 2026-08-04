from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "포럼_경영노트_신규고객_모집_전략기획서.md"
OUTPUT = ROOT / "트로이아르케_포럼_경영노트_신규고객_모집_전략기획서.docx"

FONT = "맑은 고딕"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "1F2C55"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
INPUT_FILL = "FFF4CC"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def add_inline_markdown(paragraph, text, size=11, color="000000"):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_custom_numbering(doc, num_fmt, level_text):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def widths_for(headers):
    key = "|".join(headers)
    if len(headers) == 6 and "신청 목표" in key:
        return [720, 1800, 1260, 1350, 1710, 2520]
    if len(headers) == 6 and "인터뷰 원장" in key:
        return [780, 1680, 2580, 1440, 1440, 1440]
    if len(headers) == 2 and headers[0] == "항목":
        return [2160, 7200]
    if len(headers) == 2 and headers[0] == "채널":
        return [1800, 7560]
    if len(headers) == 2 and headers[0] == "시기":
        return [1800, 7560]
    count = len(headers)
    base = CONTENT_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("TROIAREUKE PARTNERS  |  신규고객 모집 전략기획서")
    set_run_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("TROIAREUKE  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_number(p)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("TROIAREUKE PARTNERS")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("파트너스 포럼")
    set_run_font(r, size=24, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("신규고객 모집 전략기획서")
    set_run_font(r, size=15, bold=True, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    add_inline_markdown(meta, "작성 기준일  |  2026년 7월 30일", size=9.5, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    add_inline_markdown(meta, "모집 대상  |  현재 트로이아르케와 거래하지 않는 에스테틱 원장", size=9.5, color=MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(16)
    add_inline_markdown(meta, "행사 방식  |  서울·충청·대구·부산·광주 지역별 오프라인 진행", size=9.5, color=MUTED)


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    widths = widths_for(headers)

    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if len(headers) == 6 or value in {"기입", "지역 기입"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(p, value, size=8.3 if len(headers) >= 6 else 9.2)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
            elif value == "기입":
                set_cell_shading(cell, INPUT_FILL)

    set_repeat_header(table.rows[0])
    set_table_geometry(table, widths)

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def build():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    idx = 0
    title_lines_skipped = 0
    bullet_num_id = add_custom_numbering(doc, "bullet", "•")
    numbered_num_id = None
    in_numbered = False

    while idx < len(lines):
        line = lines[idx].strip()
        if line.startswith("# ") and title_lines_skipped < 1:
            title_lines_skipped += 1
            idx += 1
            continue
        if line.startswith("## ") and title_lines_skipped < 2:
            title_lines_skipped += 1
            idx += 1
            continue
        if line.startswith(("작성 기준일:", "모집 대상:", "행사 방식:")):
            idx += 1
            continue
        if not line or line == "---":
            in_numbered = False
            idx += 1
            continue
        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            if rows:
                add_table(doc, rows)
            in_numbered = False
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:], style="Heading 1")
            p.paragraph_format.keep_with_next = True
            in_numbered = False
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 2")
            p.paragraph_format.keep_with_next = True
            in_numbered = False
        elif line.startswith("#### "):
            p = doc.add_paragraph(line[5:], style="Heading 3")
            p.paragraph_format.keep_with_next = True
            in_numbered = False
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.10
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F4F6F9")
            p_pr.append(shd)
            add_inline_markdown(p, line[2:], size=10.5, color=NAVY)
        elif line.startswith("- "):
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            add_inline_markdown(p, line[2:])
            in_numbered = False
        elif re.match(r"^\d+\.\s+", line):
            if not in_numbered:
                numbered_num_id = add_custom_numbering(doc, "decimal", "%1.")
                in_numbered = True
            p = doc.add_paragraph()
            apply_num(p, numbered_num_id)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.10
            add_inline_markdown(p, line)
            in_numbered = False
        idx += 1

    doc.core_properties.title = "트로이아르케 파트너스 포럼 신규고객 모집 전략기획서"
    doc.core_properties.subject = "지역별 신규고객 모집 실행 전략"
    doc.core_properties.author = "트로이아르케"
    doc.core_properties.keywords = "트로이아르케, 파트너스 포럼, 신규고객 모집"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
